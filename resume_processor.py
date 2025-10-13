"""
Resume Processing Pipeline
=========================

This module provides functionality to extract text from resume files and
perform NLP-based information extraction for structured data.

Features:
- Text extraction from .docx, .pdf, .txt files
- OCR support for image-based PDFs
- Named Entity Recognition for resume data
- Structured information extraction
"""

import os
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import spacy
from spacy.matcher import Matcher
import pandas as pd

# File processing libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. .docx files cannot be processed.")

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PyPDF2 not installed. .pdf files cannot be processed.")

try:
    import pytesseract
    from PIL import Image
    import pdf2image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("Warning: OCR libraries not installed. Image-based PDFs cannot be processed.")

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ResumeProcessor:
    """
    Main class for processing resumes and extracting structured information.
    """
    
    def __init__(self):
        """Initialize the resume processor."""
        self.nlp = None
        self.matcher = None
        self._load_spacy_model()
        self._setup_matchers()
        
    def _load_spacy_model(self):
        """Load spaCy model for NER."""
        try:
            self.nlp = spacy.load("en_core_web_sm")
            logger.info("Loaded spaCy model: en_core_web_sm")
        except OSError:
            logger.error("spaCy model 'en_core_web_sm' not found. Please install it using: python -m spacy download en_core_web_sm")
            self.nlp = None
            
    def _setup_matchers(self):
        """Set up spaCy matchers for specific patterns."""
        if not self.nlp:
            return
            
        self.matcher = Matcher(self.nlp.vocab)
        
        # Email pattern
        email_pattern = [{"TEXT": {"REGEX": r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"}}]
        self.matcher.add("EMAIL", [email_pattern])
        
        # Phone pattern (various formats)
        phone_patterns = [
            [{"TEXT": {"REGEX": r"(\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}"}}],
            [{"TEXT": {"REGEX": r"\+?[0-9]{1,4}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}"}}]
        ]
        self.matcher.add("PHONE", phone_patterns)
        
        # LinkedIn pattern
        linkedin_pattern = [{"TEXT": {"REGEX": r"(linkedin\.com/in/|linkedin\.com/pub/)"}}]
        self.matcher.add("LINKEDIN", [linkedin_pattern])
        
    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text from various file formats.
        
        Args:
            file_path (str): Path to the resume file
            
        Returns:
            str: Extracted text content
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.txt':
                return self._extract_from_txt(file_path)
            elif extension == '.docx':
                return self._extract_from_docx(file_path)
            elif extension == '.pdf':
                return self._extract_from_pdf(file_path)
            else:
                logger.warning(f"Unsupported file format: {extension}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from .txt file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from .docx file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library is required for .docx files")
        
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file, with OCR fallback."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 library is required for .pdf files")
        
        text = ""
        
        # Try regular PDF text extraction first
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_parts = []
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                
                text = "\n".join(text_parts)
                
                # If text extraction yields very little text, try OCR
                if len(text.strip()) < 100 and OCR_AVAILABLE:
                    logger.info(f"PDF text extraction yielded little text, trying OCR for {file_path}")
                    text = self._extract_from_pdf_ocr(file_path)
                    
        except Exception as e:
            logger.warning(f"Regular PDF extraction failed for {file_path}: {str(e)}")
            if OCR_AVAILABLE:
                text = self._extract_from_pdf_ocr(file_path)
            else:
                raise e
                
        return text
    
    def _extract_from_pdf_ocr(self, file_path: Path) -> str:
        """Extract text from PDF using OCR."""
        if not OCR_AVAILABLE:
            raise ImportError("OCR libraries are required for image-based PDFs")
        
        try:
            # Convert PDF pages to images
            images = pdf2image.convert_from_path(file_path)
            text_parts = []
            
            for i, image in enumerate(images):
                # Use Tesseract OCR
                page_text = pytesseract.image_to_string(image)
                if page_text.strip():
                    text_parts.append(page_text)
                    
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"OCR extraction failed for {file_path}: {str(e)}")
            return ""
    
    def extract_entities(self, text: str) -> Dict[str, List[str]]:
        """
        Extract named entities from resume text.
        
        Args:
            text (str): Resume text content
            
        Returns:
            Dict[str, List[str]]: Extracted entities organized by type
        """
        if not self.nlp:
            logger.error("spaCy model not loaded. Cannot extract entities.")
            return {}
        
        doc = self.nlp(text)
        entities = {
            'emails': [],
            'phones': [],
            'linkedin': [],
            'skills': [],
            'companies': [],
            'education': [],
            'locations': [],
            'dates': []
        }
        
        # Extract using spaCy NER
        for ent in doc.ents:
            if ent.label_ == 'PERSON':
                # Could be candidate name, but we'll focus on other entities
                pass
            elif ent.label_ == 'ORG':
                entities['companies'].append(ent.text.strip())
            elif ent.label_ == 'GPE':
                entities['locations'].append(ent.text.strip())
            elif ent.label_ == 'DATE':
                entities['dates'].append(ent.text.strip())
        
        # Extract using custom matchers
        matches = self.matcher(doc)
        for match_id, start, end in matches:
            span = doc[start:end]
            label = self.nlp.vocab.strings[match_id]
            
            if label == 'EMAIL':
                entities['emails'].append(span.text.strip())
            elif label == 'PHONE':
                entities['phones'].append(span.text.strip())
            elif label == 'LINKEDIN':
                entities['linkedin'].append(span.text.strip())
        
        # Extract skills using keyword matching
        entities['skills'] = self._extract_skills(text)
        
        # Extract education information
        entities['education'] = self._extract_education(text)
        
        # Remove duplicates and clean up
        for key in entities:
            entities[key] = list(set([item for item in entities[key] if item.strip()]))
        
        return entities
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical skills from resume text."""
        # Common technical skills patterns
        skill_patterns = [
            r'\b(?:Java|Python|JavaScript|C\+\+|C#|PHP|Ruby|Go|Rust|Swift|Kotlin)\b',
            r'\b(?:React|Angular|Vue|Node\.?js|Express|Django|Flask|Spring|Laravel)\b',
            r'\b(?:AWS|Azure|GCP|Docker|Kubernetes|Jenkins|Git|SVN)\b',
            r'\b(?:MySQL|PostgreSQL|MongoDB|Redis|Oracle|SQL Server)\b',
            r'\b(?:Machine Learning|AI|Data Science|TensorFlow|PyTorch|Pandas|NumPy)\b',
            r'\b(?:HTML|CSS|Bootstrap|jQuery|SASS|LESS)\b',
            r'\b(?:Agile|Scrum|Kanban|DevOps|CI/CD|TDD|BDD)\b',
            r'\b(?:Project Management|Business Analysis|Product Management)\b'
        ]
        
        skills = set()
        text_lower = text.lower()
        
        for pattern in skill_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            skills.update(matches)
        
        return list(skills)
    
    def _extract_education(self, text: str) -> List[str]:
        """Extract education information from resume text."""
        education_patterns = [
            r'\b(?:Bachelor|Master|PhD|B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?|B\.?Tech|M\.?Tech|MBA|BBA)\b.*?(?:\d{4}|\d{2})',
            r'\b(?:University|College|Institute|School)\s+of\s+.*?(?:\d{4}|\d{2})',
            r'\b[A-Z][a-z]+\s+(?:University|College|Institute)\b.*?(?:\d{4}|\d{2})'
        ]
        
        education = set()
        
        for pattern in education_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            education.update(matches)
        
        return list(education)
    
    def process_resume(self, file_path: str) -> Dict:
        """
        Process a single resume file and extract all information.
        
        Args:
            file_path (str): Path to the resume file
            
        Returns:
            Dict: Complete resume information
        """
        logger.info(f"Processing resume: {file_path}")
        
        # Extract text
        text = self.extract_text_from_file(file_path)
        
        if not text.strip():
            logger.warning(f"No text extracted from {file_path}")
            return {}
        
        # Extract entities
        entities = self.extract_entities(text)
        
        # Combine results
        result = {
            'file_path': file_path,
            'file_name': Path(file_path).name,
            'text_length': len(text),
            'raw_text': text[:500] + "..." if len(text) > 500 else text,  # First 500 chars
            'entities': entities
        }
        
        return result
    
    def process_batch(self, directory_path: str, file_patterns: List[str] = None) -> List[Dict]:
        """
        Process multiple resume files from a directory.
        
        Args:
            directory_path (str): Path to directory containing resumes
            file_patterns (List[str]): File extensions to process (e.g., ['.docx', '.pdf'])
            
        Returns:
            List[Dict]: List of processed resume information
        """
        if file_patterns is None:
            file_patterns = ['.docx', '.pdf', '.txt']
        
        directory = Path(directory_path)
        results = []
        
        logger.info(f"Processing resumes from: {directory_path}")
        
        for pattern in file_patterns:
            files = list(directory.glob(f"*{pattern}")) + list(directory.glob(f"*{pattern.upper()}"))
            
            for file_path in files:
                try:
                    result = self.process_resume(str(file_path))
                    if result:
                        results.append(result)
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {str(e)}")
                    continue
        
        logger.info(f"Processed {len(results)} resumes successfully")
        return results


def main():
    """Example usage of the ResumeProcessor."""
    processor = ResumeProcessor()
    
    # Process a single file
    resume_dir = "/Users/suhasbajjuri/MyData/College FIles/Year III/SEM-5/Assignments/NLP/palaksood97/resume-dataset/versions/1/Resumes"
    
    # Test with a few files
    sample_files = list(Path(resume_dir).glob("*.docx"))[:3]
    
    for file_path in sample_files:
        print(f"\n{'='*50}")
        print(f"Processing: {file_path.name}")
        print('='*50)
        
        result = processor.process_resume(str(file_path))
        
        if result:
            print(f"Text length: {result['text_length']}")
            print(f"Preview: {result['raw_text'][:200]}...")
            
            entities = result['entities']
            print(f"\nExtracted entities:")
            for entity_type, values in entities.items():
                if values:
                    print(f"  {entity_type}: {values}")
        else:
            print("Failed to process file")


if __name__ == "__main__":
    main()
