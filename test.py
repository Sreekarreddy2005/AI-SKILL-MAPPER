import os
from pypdf import PdfReader # PdfReaderError removed here
from docx import Document
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from tempfile import TemporaryDirectory
from pathlib import Path
import json

# --- Configuration (IMPORTANT: Update these paths for your environment) ---
# NOTE: If you get 'TesseractNotFoundError', uncomment and update the tesseract_cmd path.
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# NOTE: If you get 'PDFInfoNotInstalledError', uncomment and update the Poppler path.
# POPPLER_PATH = r'C:\path\to\poppler-xx\Library\bin' 
POPPLER_PATH = None # Set to None if tesseract/poppler is in system PATH

# --------------------------------------------------------------------------


def extract_text_from_file(file_path: str) -> str:
    """
    Attempts to extract text from .pdf, .docx, or .txt files, 
    using OCR as a fallback for image-based PDFs.
    """
    file_extension = Path(file_path).suffix.lower()
    full_text = ""

    if not os.path.exists(file_path):
        return f"Error: File not found at path: {file_path}"

    if file_extension == '.txt':
        # --- Handle .txt files (Simplest Case) ---
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                full_text = f.read()
        except Exception as e:
            print(f"Error reading .txt file: {e}")
            full_text = ""
            
    elif file_extension == '.docx':
        # --- Handle .docx files ---
        try:
            document = Document(file_path)
            full_text = '\n'.join([p.text for p in document.paragraphs])
        except Exception as e:
            print(f"Error reading .docx file: {e}")
            full_text = ""

    elif file_extension == '.pdf':
        # --- Handle .pdf files (Primary Logic: PyPDF -> OCR) ---
        
        # 1. Attempt Native Text Extraction (for digitally-born PDFs)
        try:
            reader = PdfReader(file_path)
            page_texts = [page.extract_text() for page in reader.pages if page.extract_text()]
            
            # If we get *some* text, use it. This is faster and generally cleaner.
            if "".join(page_texts).strip():
                return "\n".join(page_texts)

        except Exception: # Changed from PdfReaderError for better compatibility
             # This often happens with password-protected or corrupted files.
            print("PyPDF failed to read or extract text. Proceeding to OCR...")
        except Exception as e:
            # Catch other non-OCR related PDF errors
            print(f"Non-OCR PDF extraction failed: {e}. Proceeding to OCR...")


        # 2. OCR Fallback (for image-based or difficult PDFs)
        print(f"[{Path(file_path).name}] Falling back to OCR...")
        try:
            with TemporaryDirectory() as tempdir:
                # Convert PDF pages to images
                images = convert_from_path(
                    file_path, 
                    poppler_path=POPPLER_PATH,
                    dpi=300 # Higher DPI for better OCR results
                )

                ocr_texts = []
                for image in images:
                    # Use pytesseract to extract text from each image
                    text = pytesseract.image_to_string(image)
                    ocr_texts.append(text)
                
                full_text = "\n".join(ocr_texts)

        except Exception as e:
            print(f"OCR failed for PDF {file_path}. Error: {e}")
            full_text = ""
            
    else:
        full_text = f"Unsupported file format: {file_extension}"

    return full_text.strip()
    

def create_dummy_files():
    """Helper to create dummy files for basic testing."""
    test_dir = "test_resumes"
    if not os.path.exists(test_dir):
        os.makedirs(test_dir)

    # 1. Simple TXT file
    with open(os.path.join(test_dir, "test_simple.txt"), "w") as f:
        f.write("Candidate: Alice Johnson. Email: alice@corp.com\nSkills: Python, SQL.")

    # 2. Simple DOCX file
    try:
        from docx import Document as DocxCreator
        document = DocxCreator()
        document.add_paragraph("Work Experience\nSenior Analyst at Acme Corp (2020-Present)")
        document.add_paragraph("Education: MS Data Science, 2018")
        document.save(os.path.join(test_dir, "test_simple.docx"))
    except ImportError:
        print("Note: Skipping DOCX creation. Ensure 'python-docx' is installed for DOCX testing.")
    except Exception as e:
        print(f"Note: Skipping DOCX creation due to error: {e}")
        
    # NOTE: Cannot easily create a dummy PDF file without external tools.
    # You MUST use the PDF files downloaded from Kaggle for PDF/OCR testing.


def process_resume_dataset():
    """Process the actual resume dataset and extract information."""
    print("--- 📄 Resume Dataset Processing ---")
    
    # Path to the actual resume dataset
    resume_dir = Path("palaksood97/resume-dataset/versions/1/Resumes")
    
    if not resume_dir.exists():
        print(f"❌ Resume dataset not found at: {resume_dir}")
        print("Please run datasets.py first to download the dataset.")
        return
    
    # Get all .docx files
    resume_files = list(resume_dir.glob("*.docx")) + list(resume_dir.glob("*.DOCX"))
    
    if not resume_files:
        print(f"❌ No resume files found in: {resume_dir}")
        return
    
    print(f"✅ Found {len(resume_files)} resume files")
    
    # Process first 5 resumes for testing
    processed_resumes = []
    
    for i, resume_file in enumerate(resume_files[:5]):
        print(f"\n[Processing {i+1}/5] {resume_file.name}")
        
        try:
            text = extract_text_from_file(str(resume_file))
            
            if text and not text.startswith("Error:"):
                resume_data = {
                    "file_name": resume_file.name,
                    "text_length": len(text),
                    "text_preview": text[:300] + "..." if len(text) > 300 else text,
                    "full_text": text
                }
                processed_resumes.append(resume_data)
                print(f"✅ Successfully processed - {len(text)} characters")
            else:
                print(f"❌ Failed to extract text: {text[:100]}...")
                
        except Exception as e:
            print(f"❌ Error processing {resume_file.name}: {str(e)}")
    
    # Save results
    if processed_resumes:
        output_file = "processed_resumes.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(processed_resumes, f, indent=2, ensure_ascii=False)
        
        print(f"\n💾 Processed {len(processed_resumes)} resumes saved to: {output_file}")
        
        # Show summary
        print(f"\n📊 Summary:")
        print(f"  Total resumes processed: {len(processed_resumes)}")
        avg_length = sum(r['text_length'] for r in processed_resumes) / len(processed_resumes)
        print(f"  Average text length: {avg_length:.0f} characters")
        
        # Show sample output
        print(f"\n📄 Sample Output (first resume):")
        print("-" * 50)
        print(processed_resumes[0]['text_preview'])
        print("-" * 50)
    
    return processed_resumes


if __name__ == "__main__":
    # Test with dummy files first
    create_dummy_files()
    
    # --- Paths for the generated dummy files ---
    path_txt = "test_resumes/test_simple.txt"
    path_docx = "test_resumes/test_simple.docx"
    path_unsupported = "test_resumes/image.png"
    
    print("--- 📄 Resume Text Extraction Test Suite ---")
    
    # Test 1: .TXT File
    print(f"\n[Test 1] Processing {path_txt}:")
    print(extract_text_from_file(path_txt))
    
    # Test 2: .DOCX File
    if os.path.exists(path_docx):
        print(f"\n[Test 2] Processing {path_docx}:")
        print(extract_text_from_file(path_docx))
    
    # Test 3: Unsupported File
    print(f"\n[Test 3] Processing Unsupported File (.png):")
    print(extract_text_from_file(path_unsupported))
    
    print("\n" + "="*60)
    
    # Process actual resume dataset
    process_resume_dataset()